from copy import deepcopy
from docx import Document

resume_template = '/Users/alexandercarnevale/repos/tailor_resume/assets/base_resume_rev_ops.docx'
output_doc = '/Users/alexandercarnevale/repos/tailor_resume/assets/AMC_Generic.docx'


# -------------------------
# PLACEHOLDERS IN DOC
# -------------------------
CONSULTING_BULLETS_PLACEHOLDER = "[CONSULTING_BULLETS_HERE]"
ADP_BULLETS_PLACEHOLDER = "[ADP_BULLETS_HERE]"

CONSULTING_TITLE_PLACEHOLDER = "[CONSULTING_TITLE_HERE]"
CORE_COMPETENCIES_PLACEHOLDER = "[CORE_COMPETENCIES_HERE]"
SUMMARY_PLACEHOLDER = "[SUMMARY_HERE]"
TARGET_TITLE_PLACEHOLDER = "[TARGET_TITLE]"
REVENUE_IMPACT_BULLETS_PLACEHOLDER = "[REVENUE_IMPACT_BULLETS_HERE]"
TOOLS_AND_SYSTEMS_PLACEHOLDER = "[TOOLS_SYSTEMS_HERE]"


# -------------------------
# INPUT DATA
# -------------------------
TARGET_TITLE = "Revenue Operations Manager"
CONSULTING_TITLE = "Senior Consultant"

SUMMARY = "Revenue Operations Manager with experience owning pipeline, forecasting, and CRM systems across B2B SaaS environments. Specialized in HubSpot, lifecycle management, and revenue reporting, with a track record of improving conversion, increasing forecast accuracy, and enabling data-driven decision-making across Sales, Marketing, and Customer Success."

REVENUE_IMPACT_BULLETS = """•	Increased lead generation by 50% through lifecycle optimization and funnel design
•	Improved forecast accuracy by ~10% by implementing pipeline governance and validation frameworks
•	Supported revenue growth from ~$1M to $20M+ by improving forecasting, reporting, and operational alignment
"""

CORE_COMPETENCIES = (
    "HubSpot & CRM Management | Pipeline & Forecasting | Revenue Reporting & Dashboards | "
    "GTM Process Optimization | Lifecycle Management | Data Integrity & Governance | Cross-Functional Operations"
)

TOOLS_AND_SYSTEMS = "Salesforce CRM (reporting & workflow optimization) | Zoho CRM | Power BI (dashboard development) | Looker | Tableau | Excel (advanced modeling) | Commission & incentive modeling"

RAW_CONSULTING_BULLETS = """
• Own end-to-end revenue operations across lead management, pipeline progression, forecasting, and customer lifecycle, improving visibility and execution across Sales, Marketing, and Customer Success  
• Build and maintain HubSpot infrastructure (lifecycle stages, pipelines, automation, and reporting), ensuring accurate data, scalable processes, and alignment across revenue teams  
• Increase lead generation by 50% and improve forecast accuracy by ~10% by standardizing lifecycle definitions, pipeline governance, and data validation practices
• Partner with Sales, Marketing, and Customer Success to improve funnel performance, reduce cycle time, and increase conversion across the lead-to-customer lifecycle
• Establish and enforce data governance and CRM best practices, ensuring data accuracy, consistency, and reliability across reporting and forecasting
• Identify and address pipeline bottlenecks through ongoing analysis of funnel metrics, driving improvements in deal velocity, conversion rates, and overall revenue efficiency
• Support pipeline reviews and forecasting cadence with Sales leadership, providing visibility into performance, risks, and opportunities
"""

RAW_ADP_BULLETS = """
• Supported revenue growth from ~$1M to $20M+ by improving forecasting, renewal visibility, and revenue reporting across enterprise accounts
• Built and maintained Salesforce and Looker dashboards tracking pipeline, bookings, renewals, and commission exposure, enabling better visibility into revenue performance and risk
• Partnered with Sales, Finance, and Delivery to improve lead-to-cash processes, aligning pipeline activity, revenue recognition, and forecasting accuracy
• Developed reporting and analysis across 20+ enterprise clients and 5,000+ users, identifying revenue drivers, margin risks, and expansion opportunities
• Improved data consistency and reporting accuracy by standardizing processes and definitions across systems and cross-functional teams
• Designed and scaled operational processes supporting team growth from 4 to 20 employees while maintaining visibility, accountability, and execution discipline
• Led post-merger integration of systems, reporting, and workflows, ensuring continuity across customer lifecycle, revenue tracking, and cross-functional operations
• Supported sales compensation and quota tracking through reporting and analysis of bookings, renewals, and performance metrics
"""


def parse_bullets(raw_text: str) -> list[str]:
    bullets = []

    for line in raw_text.strip().splitlines():
        line = line.strip()
        if not line:
            continue

        for prefix in ("•", "-", "*"):
            if line.startswith(prefix):
                line = line[len(prefix):].strip()
                break

        if line:
            bullets.append(line)

    return bullets


def clone_paragraph_after(paragraph):
    new_p = deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    return new_p


def set_paragraph_text_preserve_format(paragraph, text: str):
    if not paragraph.runs:
        paragraph.add_run(text)
        return

    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def replace_placeholder_paragraph_with_text(doc, placeholder: str, text: str):
    for para in doc.paragraphs:
        if para.text.strip() == placeholder:
            set_paragraph_text_preserve_format(para, text)
            return

    print(f"Warning: placeholder not found -> {placeholder}")


def replace_placeholder_with_bullets(doc, placeholder: str, bullet_items: list[str]):
    if not bullet_items:
        return

    target_para = None
    for para in doc.paragraphs:
        if para.text.strip() == placeholder:
            target_para = para
            break

    if target_para is None:
        print(f"Warning: placeholder not found -> {placeholder}")
        return

    # first bullet uses original paragraph formatting
    set_paragraph_text_preserve_format(target_para, bullet_items[0])
    current_para = target_para

    # each cloned paragraph preserves the template paragraph formatting
    for bullet_text in bullet_items[1:]:
        new_p_xml = clone_paragraph_after(current_para)
        new_para = current_para._parent.paragraphs[
            current_para._parent._element.index(new_p_xml)
        ]
        set_paragraph_text_preserve_format(new_para, bullet_text)
        current_para = new_para


def main():
    doc = Document(resume_template)

    replace_placeholder_with_bullets(
        doc,
        CONSULTING_BULLETS_PLACEHOLDER,
        parse_bullets(RAW_CONSULTING_BULLETS),
    )

    replace_placeholder_with_bullets(
        doc,
        ADP_BULLETS_PLACEHOLDER,
        parse_bullets(RAW_ADP_BULLETS),
    )

    replace_placeholder_with_bullets(
        doc,
        REVENUE_IMPACT_BULLETS_PLACEHOLDER,
        parse_bullets(REVENUE_IMPACT_BULLETS),
    )

    replace_placeholder_paragraph_with_text(doc, CONSULTING_TITLE_PLACEHOLDER, CONSULTING_TITLE.upper())
    replace_placeholder_paragraph_with_text(doc, CORE_COMPETENCIES_PLACEHOLDER, CORE_COMPETENCIES)
    replace_placeholder_paragraph_with_text(doc, SUMMARY_PLACEHOLDER, SUMMARY)
    replace_placeholder_paragraph_with_text(doc, TARGET_TITLE_PLACEHOLDER, TARGET_TITLE)
    replace_placeholder_paragraph_with_text(doc, TOOLS_AND_SYSTEMS_PLACEHOLDER, TOOLS_AND_SYSTEMS)

    doc.save(output_doc)
    print("Done.")


if __name__ == "__main__":
    main()